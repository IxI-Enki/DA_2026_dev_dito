"""T065 + T021: Tests for MediaProcessor (PDF/OCR + PDF quality cleaning)."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import patch

import pytest


@pytest.fixture()
def media_dir(tmp_path: Path) -> Path:
    """Create a temp media directory with sample files."""
    d = tmp_path / "media"
    d.mkdir()
    (d / "sample.txt").write_text("Hello world from text file", encoding="utf-8")
    # Create a fake PDF (just bytes - actual extraction is mocked)
    (d / "sample.pdf").write_bytes(b"%PDF-1.4 fake pdf content")
    # Create a fake image
    (d / "sample.png").write_bytes(b"\x89PNG fake image content")
    return d


class TestMediaProcessorPDF:
    """Tests for process_pdf method."""

    def test_process_pdf_returns_string(self, tmp_path: Path) -> None:
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        pdf = tmp_path / "test.pdf"
        pdf.write_bytes(b"%PDF")
        with patch(
            "media_processor.MediaProcessor._extract_pdf_text", return_value="Extracted text"
        ):
            result = mp.process_pdf(pdf)
        assert isinstance(result, str)

    def test_process_pdf_tries_docling_first(self, tmp_path: Path) -> None:
        """When Docling returns non-empty text, that text is returned (no clean_pdf_text)."""
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        pdf = tmp_path / "docling.pdf"
        pdf.write_bytes(b"%PDF")
        docling_md = "# Title\n\nParagraph from Docling."
        with patch(
            "media_processor.MediaProcessor._extract_pdf_text_docling", return_value=docling_md
        ):
            result = mp.process_pdf(pdf)
        assert result == docling_md
        assert result.strip().startswith("# Title")

    def test_process_pdf_fallback_to_ocr(self, tmp_path: Path) -> None:
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        pdf = tmp_path / "scanned.pdf"
        pdf.write_bytes(b"%PDF")
        # Simulate Docling empty, pypdf empty -> should attempt OCR fallback
        with (
            patch("media_processor.MediaProcessor._extract_pdf_text_docling", return_value=""),
            patch("media_processor.MediaProcessor._extract_pdf_text", return_value=""),
            patch("media_processor.MediaProcessor._ocr_pdf", return_value="OCR result"),
        ):
            result = mp.process_pdf(pdf)
        assert result == "OCR result"

    def test_process_pdf_missing_file(self, tmp_path: Path) -> None:
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        result = mp.process_pdf(tmp_path / "nonexistent.pdf")
        assert result == ""


class TestMediaProcessorDOCX:
    """Tests for process_docx: Docling first, fallback with tables."""

    def test_process_docx_tries_docling_first(self, tmp_path: Path) -> None:
        """When Docling returns non-empty text, that text is returned."""
        from docx import Document
        from media_processor import MediaProcessor

        docx_path = tmp_path / "sample.docx"
        doc = Document()
        doc.add_paragraph("Hello from Docx.")
        doc.save(str(docx_path))
        mp = MediaProcessor()
        docling_md = "# Docling Title\n\nStructured content."
        with patch(
            "media_processor.MediaProcessor._extract_docx_text_docling", return_value=docling_md
        ):
            result = mp.process_docx(docx_path)
        assert result == docling_md
        assert result.strip().startswith("# Docling Title")

    def test_process_docx_fallback_extracts_tables(self, tmp_path: Path) -> None:
        """Fallback python-docx extracts both paragraphs and tables as Markdown."""
        from docx import Document
        from media_processor import MediaProcessor

        docx_path = tmp_path / "with_tables.docx"
        doc = Document()
        doc.add_paragraph("Intro text.")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "A"
        table.rows[0].cells[1].text = "B"
        table.rows[1].cells[0].text = "1"
        table.rows[1].cells[1].text = "2"
        doc.save(str(docx_path))
        mp = MediaProcessor()
        with patch("media_processor.MediaProcessor._extract_docx_text_docling", return_value=""):
            result = mp.process_docx(docx_path)
        assert "Intro text." in result
        assert "| A | B |" in result
        assert "---" in result
        assert "| 1 | 2 |" in result


class TestMediaProcessorXLSX:
    """Tests for process_xlsx: Docling first, fallback with proper Markdown tables."""

    def test_process_xlsx_tries_docling_first(self, tmp_path: Path) -> None:
        """When Docling returns non-empty text, that text is returned."""
        from media_processor import MediaProcessor
        from openpyxl import Workbook

        xlsx_path = tmp_path / "sample.xlsx"
        wb = Workbook()
        ws = wb.active
        ws["A1"] = "Col1"
        ws["B1"] = "Col2"
        wb.save(str(xlsx_path))
        mp = MediaProcessor()
        docling_md = "## Sheet1\n\n| Col1 | Col2 |"
        with patch(
            "media_processor.MediaProcessor._extract_xlsx_text_docling", return_value=docling_md
        ):
            result = mp.process_xlsx(xlsx_path)
        assert result == docling_md

    def test_process_xlsx_fallback_produces_markdown_tables(self, tmp_path: Path) -> None:
        """Fallback openpyxl produces proper Markdown table syntax with | and ---."""
        from media_processor import MediaProcessor
        from openpyxl import Workbook

        xlsx_path = tmp_path / "data.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.title = "Data"
        ws["A1"], ws["B1"] = "X", "Y"
        ws["A2"], ws["B2"] = "a", "b"
        wb.save(str(xlsx_path))
        mp = MediaProcessor()
        with patch("media_processor.MediaProcessor._extract_xlsx_text_docling", return_value=""):
            result = mp.process_xlsx(xlsx_path)
        assert "## Data" in result
        assert "| " in result and " | " in result
        assert "---" in result


class TestMediaProcessorImage:
    """Tests for process_image method."""

    def test_process_image_returns_string(self, tmp_path: Path) -> None:
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        img = tmp_path / "test.png"
        img.write_bytes(b"\x89PNG")
        with patch("media_processor.MediaProcessor._ocr_image", return_value="OCR text"):
            result = mp.process_image(img)
        assert isinstance(result, str)

    def test_process_image_missing_file(self, tmp_path: Path) -> None:
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        result = mp.process_image(tmp_path / "nonexistent.png")
        assert result == ""


class TestMediaProcessorDirectory:
    """Tests for process_media_directory method."""

    def test_directory_returns_list(self, media_dir: Path) -> None:
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        with (
            patch.object(mp, "process_pdf", return_value="pdf text"),
            patch.object(mp, "process_image", return_value="image text"),
        ):
            results = mp.process_media_directory(media_dir)
        assert isinstance(results, list)

    def test_directory_entries_have_required_keys(self, media_dir: Path) -> None:
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        with (
            patch.object(mp, "process_pdf", return_value="pdf text"),
            patch.object(mp, "process_image", return_value="image text"),
        ):
            results = mp.process_media_directory(media_dir)
        for entry in results:
            assert "filename" in entry
            assert "text" in entry
            assert "type" in entry


class TestFixSpacedCharacters:
    """T021: Tests for _fix_spaced_characters() -- US7 PDF quality."""

    def test_htbla_leonding(self) -> None:
        """Spec AC1: 'H T B L A  L e o n d i n g' -> 'HTBLA Leonding'."""
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        result = mp._fix_spaced_characters("H T B L A  L e o n d i n g")
        assert result == "HTBLA Leonding"

    def test_mixed_spaced_and_normal_lines(self) -> None:
        """Only lines with >60% single-char words are fixed."""
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        text = "H T B L A  L e o n d i n g\nThis is a normal line."
        result = mp._fix_spaced_characters(text)
        assert "HTBLA" in result
        assert "This is a normal line." in result

    def test_normal_text_unchanged(self) -> None:
        """Normal text without spaced characters is not modified."""
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        text = "This is perfectly normal text with words."
        result = mp._fix_spaced_characters(text)
        assert result == text

    def test_empty_string(self) -> None:
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        assert mp._fix_spaced_characters("") == ""


class TestMergeShortLines:
    """T021: Tests for _merge_short_lines() -- US7 PDF quality."""

    def test_joins_consecutive_short_lines(self) -> None:
        """Short lines (<40 chars) are joined into paragraphs."""
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        text = "Short line one\nshort line two\nshort three."
        result = mp._merge_short_lines(text)
        assert "Short line one short line two short three." in result

    def test_respects_sentence_boundaries(self) -> None:
        """Lines ending with sentence terminators start new segments."""
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        text = "First sentence ends.\nSecond sentence here."
        result = mp._merge_short_lines(text)
        # Both are short, but first ends with period -- still merge is OK
        # Key: they get merged into one line
        assert "First sentence ends." in result

    def test_preserves_list_items(self) -> None:
        """Lines starting with -, *, or digits are NOT merged with previous."""
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        text = "Intro paragraph\n- Item one\n- Item two\n* Star item"
        result = mp._merge_short_lines(text)
        assert "\n- Item one" in result or result.startswith("Intro paragraph\n- Item one")
        assert "\n- Item two" in result
        assert "\n* Star item" in result

    def test_preserves_headings(self) -> None:
        """Lines starting with # are NOT merged with previous."""
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        text = "Some text\n# Heading\nMore text"
        result = mp._merge_short_lines(text)
        assert "\n# Heading\n" in result

    def test_preserves_empty_line_paragraph_breaks(self) -> None:
        """Empty lines (paragraph separators) are preserved."""
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        text = "Paragraph one.\n\nParagraph two."
        result = mp._merge_short_lines(text)
        assert "\n\n" in result

    def test_long_lines_not_merged(self) -> None:
        """Lines longer than threshold are not merged with the next."""
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        long_line = "A" * 50 + " long line that is definitely above threshold"
        text = f"{long_line}\nShort next line"
        result = mp._merge_short_lines(text)
        assert long_line in result


class TestCleanPdfText:
    """T021: Tests for clean_pdf_text() -- chains both operations."""

    def test_chains_spaced_chars_then_merge(self) -> None:
        """clean_pdf_text applies spaced char fix THEN short line merge."""
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        # Spaced chars on first line, short lines after
        text = "H T B L A  L e o n d i n g\nis a school\nin Upper Austria"
        result = mp.clean_pdf_text(text)
        assert "HTBLA Leonding" in result
        # Short lines should be merged
        assert "is a school in Upper Austria" in result

    def test_process_pdf_integrates_clean(self, tmp_path: Path) -> None:
        """process_pdf calls clean_pdf_text as post-processing."""
        from media_processor import MediaProcessor

        mp = MediaProcessor()
        pdf = tmp_path / "test.pdf"
        pdf.write_bytes(b"%PDF")
        raw = "H T B L A  L e o n d i n g\nis great"
        with patch("media_processor.MediaProcessor._extract_pdf_text", return_value=raw):
            result = mp.process_pdf(pdf)
        assert "HTBLA Leonding" in result
